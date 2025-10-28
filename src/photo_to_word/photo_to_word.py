import tkinter as tk
from tkinter import filedialog, ttk, messagebox
import os
from docx import Document
from docx.shared import Inches

#使用照片貼入 Word 的功能
class PhotoToWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("照片貼入Word工具")

        # 建立主框架
        self.main_frame = ttk.Frame(root, padding="10")
        self.main_frame.pack(fill=tk.BOTH, expand=True)

        # 選擇照片按鈕
        ttk.Button(self.main_frame, text="選擇照片", command=self.select_photos).pack(pady=5)

        # 已選照片數量顯示
        self.photo_count_var = tk.StringVar(value="尚未選擇照片")
        ttk.Label(self.main_frame, textvariable=self.photo_count_var).pack(pady=5)

        # 格式設定
        format_frame = ttk.LabelFrame(self.main_frame, text="表格格式設定", padding="10")
        format_frame.pack(fill=tk.X, pady=10)

        # 列數設定
        ttk.Label(format_frame, text="列數:").grid(row=0, column=0, padx=5)
        self.rows_var = tk.StringVar(value="2")
        self.rows_entry = ttk.Entry(format_frame, textvariable=self.rows_var, width=10)
        self.rows_entry.grid(row=0, column=1, padx=5)

        # 欄數設定
        ttk.Label(format_frame, text="欄數:").grid(row=0, column=2, padx=5)
        self.cols_var = tk.StringVar(value="2")
        self.cols_entry = ttk.Entry(format_frame, textvariable=self.cols_var, width=10)
        self.cols_entry.grid(row=0, column=3, padx=5)

        # 圖片寬度設定
        ttk.Label(format_frame, text="圖片寬度(英吋):").grid(row=1, column=0, columnspan=2, padx=5, pady=5)
        self.width_var = tk.StringVar(value="3.5")
        self.width_entry = ttk.Entry(format_frame, textvariable=self.width_var, width=10)
        self.width_entry.grid(row=1, column=2, padx=5, pady=5)

        # 產生按鈕
        ttk.Button(self.main_frame, text="產生Word文件", command=self.create_word_doc).pack(pady=10)

        # 儲存選擇的照片路徑
        self.selected_photos = []

    def select_photos(self):
        files = filedialog.askopenfilenames(
            title='選擇照片',
            filetypes=[
                ('圖片檔', '*.jpg;*.jpeg;*.png'),
                ('所有檔案', '*.*')
            ]
        )

        if files:
            self.selected_photos = list(files)
            self.photo_count_var.set(f"已選擇 {len(self.selected_photos)} 張照片")

    def create_word_doc(self):
        if not self.selected_photos:
            messagebox.showwarning("警告", "請先選擇照片")
            return

        try:
            rows = int(self.rows_var.get())
            cols = int(self.cols_var.get())
            width = float(self.width_var.get())
        except ValueError:
            messagebox.showerror("錯誤", "請輸入有效的數字")
            return

        # 選擇儲存位置
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx")],
            initialfile="照片集.docx"
        )

        if not save_path:
            return

        try:
            doc = Document()

            # 設定頁面邊界
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            # 計算需要的表格數量
            images_per_table = rows * cols
            total_images = len(self.selected_photos)
            tables_needed = (total_images + images_per_table - 1) // images_per_table

            for table_num in range(tables_needed):
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'

                for i in range(images_per_table):
                    image_index = table_num * images_per_table + i
                    if image_index >= total_images:
                        break

                    row = i // cols
                    col = i % cols
                    cell = table.cell(row, col)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()

                    try:
                        run.add_picture(self.selected_photos[image_index], width=Inches(width))
                    except Exception as e:
                        print(f"插入圖片失敗: {self.selected_photos[image_index]}")

                if table_num < tables_needed - 1:
                    doc.add_page_break()

            doc.save(save_path)
            messagebox.showinfo("成功", f"Word文件已儲存至:\n{save_path}")

        except Exception as e:
            messagebox.showerror("錯誤", f"建立Word文件時發生錯誤: {str(e)}")


def main():
    root = tk.Tk()
    app = PhotoToWordApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()